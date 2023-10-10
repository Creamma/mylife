import pandas as pd  
from docx import Document  
from collections import Counter  
from datetime import datetime, timedelta  
  
# 读取Excel文件  
df = pd.read_excel('myLife.xlsx', sheet_name='Scene')  
  
# 确保日期列为日期格式  
df['日期'] = pd.to_datetime(df['日期'])  
  
# 获取当前日期并筛选前7天的活动数据  
current_date = datetime.now()  
start_date = current_date - timedelta(days=7)  
end_date = current_date  
df_7days = df[(df['日期'] >= start_date) & (df['日期'] <= end_date)]  
  
# 创建新的Word文档  
doc = Document()  
  
# 将统计结果添加到Word文档的最上面  
doc.add_paragraph('以下是本周的天气统计：')  
  
# 统计这一周的最高温度，最低温度，平均温度  
max_temp = df_7days['最高温度'].max()  
min_temp = df_7days['最底温度'].min()  
avg_temp = df_7days['最高温度'].mean()  
doc.add_paragraph(f'最高温度: {max_temp}, 最底温度: {min_temp}, 平均温度: {avg_temp}')  
  
# 统计一周的天气列表出现最多次数的天气  
weather_counts = Counter(df_7days['天气'])  
most_common_weather = weather_counts.most_common(1)[0][0]  
doc.add_paragraph(f'出现最多次数的天气: {most_common_weather}')  
  
# 将Scene表的内容添加到Word文档  
doc.add_paragraph('以下是本周的Scene表内容：')  
for index, row in df_7days.iterrows():  
    doc.add_paragraph(f'日期: {row["日期"]}, 城市: {row["城市"]}, 地点: {row["地点"]}, 天气: {row["天气"]}, 最高温度: {row["最高温度"]}, 最底温度: {row["最底温度"]}, 是否工作日: {row["是否工作日"]}')  
  
# 以特定格式命名并保存Word文档  
doc_name = f'{start_date.year}-{start_date.month}-{start_date.day}_to_{end_date.year}-{end_date.month}-{end_date.day}_7days_weather.docx'  
doc.save(doc_name)